import pandas as pd
from docx import Document
from docx.table import Table
from docx2pdf import convert
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Inches, Pt


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def add_table_from_dataframe(document, paragraph_identifier, dataframe):
    # Find the paragraph after the paragraph that has the text `paragraph_identifier`
    table_paragraph_index = -1
    for i, p in enumerate(document.paragraphs):
        if p.text == paragraph_identifier:
            table_paragraph_index = i + 1
            break

    if table_paragraph_index == -1:
        raise ValueError(f"Paragraph identifier '{paragraph_identifier}' not found in document.")

    table_paragraph = document.paragraphs[table_paragraph_index]

    # Create the table with the number of columns equal to the DataFrame columns
    table = document.add_table(rows=1, cols=len(dataframe.columns), style='Grid Table 4 Accent 1')

    # Add the header row
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(dataframe.columns):
        hdr_cells[i].text = str(column)
        # Optional: Customize header formatting
        for paragraph in hdr_cells[i].paragraphs:
            run = paragraph.runs[0]
            run.font.size = Pt(12)
            run.bold = True

    # Add the data rows
    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
            # Optional: Customize cell formatting
            for paragraph in row_cells[i].paragraphs:
                run = paragraph.runs[0]
                run.font.size = Pt(12)

    # Move the table to the correct position
    move_table_after(table, table_paragraph)


def fill_report_template(template_path, output_path, data, custom_tables):
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                if type(value) == float:
                    value = round(value, 2)
                paragraph.text = paragraph.text.replace(key, str(value))

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        if type(value) == float:
                            value = round(value, 2)
                        cell.text = cell.text.replace(key, str(value))

    for table_identifier, dataframe in custom_tables.items():
        add_table_from_dataframe(doc, table_identifier, dataframe)


    # export to word and pdf file
    doc.save(output_path)
    pdf_output_path = output_path.replace('.docx', '.pdf')
    convert(output_path, pdf_output_path)